"""Эталонное решение (reference_solution_url) должно резолвиться и
обрабатываться ОДИНАКОВО для ручной кнопки "Проверить ИИ" и автопроверки по
дедлайну — и включать не только текст, но и встроенные изображения (docx/
pptx), как и для сдачи студента. Регресс на два найденных при аудите бага:

1. crud.resolve_reference_solution_urls: раньше автопроверка по дедлайну
   игнорировала submission.variant_number и всегда брала
   assignment.reference_solution_url — для заданий с вариантами это либо
   пусто, либо чужой эталон.
2. _resolve_reference_material: эталон учителя читался только как текст,
   картинки внутри reference-файла игнорировались."""
import asyncio
import io
import json as _json

import pytest
from docx import Document
from docx.shared import Inches
from PIL import Image

from crud import assignments as crud_assignments
from crud import classes as crud_classes
from models import Submission
from services import ai_grader
from services import deadline_checker
from tests.conftest import make_user, auth_headers
from tests.test_lecture_context import _CapturingClient


def _png_bytes(color: tuple) -> bytes:
    buf = io.BytesIO()
    Image.new("RGB", (20, 20), color=color).save(buf, format="PNG")
    return buf.getvalue()


def _make_assignment(db, teacher, max_score=100, reference_solution_url=None):
    cls = crud_classes.create_class(db, "Bio", None, created_by=teacher.id)
    assignment = crud_assignments.create_assignment(
        db=db, class_id=cls.id, title="HW", description=None,
        criteria=[{"name": "полнота", "weight": max_score}], max_score=max_score,
        deadline=None, created_by=teacher.id,
        reference_solution_url=reference_solution_url,
    )
    return cls, assignment


# ── crud.resolve_reference_solution_urls ────────────────────────────────────

def test_resolve_reference_no_variant_uses_assignment_level(db_session):
    teacher = make_user(db_session, role="teacher")
    _, assignment = _make_assignment(db_session, teacher, reference_solution_url="http://x/assignment.pdf")
    sub = Submission(assignment_id=assignment.id, student_id=teacher.id, variant_number=None)
    urls = crud_assignments.resolve_reference_solution_urls(db_session, assignment, sub)
    assert urls == ["http://x/assignment.pdf"]


def test_resolve_reference_with_variant_uses_variant_not_assignment(db_session):
    """Регресс: раньше автопроверка ВСЕГДА брала assignment.reference_solution_url,
    даже если сдача привязана к варианту с собственным (другим) эталоном."""
    teacher = make_user(db_session, role="teacher")
    _, assignment = _make_assignment(
        db_session, teacher, reference_solution_url="http://x/wrong-assignment-level.pdf"
    )
    crud_classes.add_variant(db_session, assignment.id, 1, "http://x/variant1.pdf")
    crud_classes.add_variant(db_session, assignment.id, 2, "http://x/variant2.pdf")

    sub1 = Submission(assignment_id=assignment.id, student_id=teacher.id, variant_number=1)
    sub2 = Submission(assignment_id=assignment.id, student_id=teacher.id, variant_number=2)

    assert crud_assignments.resolve_reference_solution_urls(db_session, assignment, sub1) == ["http://x/variant1.pdf"]
    assert crud_assignments.resolve_reference_solution_urls(db_session, assignment, sub2) == ["http://x/variant2.pdf"]


def test_resolve_reference_variant_not_found_falls_back_to_assignment(db_session):
    teacher = make_user(db_session, role="teacher")
    _, assignment = _make_assignment(db_session, teacher, reference_solution_url="http://x/fallback.pdf")
    sub = Submission(assignment_id=assignment.id, student_id=teacher.id, variant_number=99)
    urls = crud_assignments.resolve_reference_solution_urls(db_session, assignment, sub)
    assert urls == ["http://x/fallback.pdf"]


def test_resolve_reference_parses_json_array():
    urls = crud_assignments._parse_reference_urls('["http://x/a.pdf", "http://x/b.pdf"]')
    assert urls == ["http://x/a.pdf", "http://x/b.pdf"]


def test_resolve_reference_empty_when_no_reference_anywhere(db_session):
    teacher = make_user(db_session, role="teacher")
    _, assignment = _make_assignment(db_session, teacher, reference_solution_url=None)
    sub = Submission(assignment_id=assignment.id, student_id=teacher.id, variant_number=None)
    assert crud_assignments.resolve_reference_solution_urls(db_session, assignment, sub) == []


# ── _resolve_reference_material: текст + встроенные картинки эталона ────────

class _FakeGetClient:
    def __init__(self, content: bytes, content_type: str = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"):
        self._content = content
        self._content_type = content_type

    def __call__(self, *a, **k):
        return self

    async def __aenter__(self):
        return self

    async def __aexit__(self, *a):
        return False

    async def get(self, *a, **k):
        class _R:
            is_success = True
            status_code = 200
            content = self._content
            headers = {"content-type": self._content_type}
        return _R()


def _docx_with_image(text: str, image: bytes) -> bytes:
    doc = Document()
    doc.add_paragraph(text)
    doc.add_paragraph().add_run().add_picture(io.BytesIO(image), width=Inches(1))
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_resolve_reference_material_extracts_text_and_images(monkeypatch):
    ref_image = _png_bytes((1, 2, 3))
    data = _docx_with_image("Эталонный ответ учителя.", ref_image)
    monkeypatch.setattr(ai_grader.httpx, "AsyncClient", _FakeGetClient(data))

    text, images = asyncio.run(ai_grader._resolve_reference_material(["http://localhost:8000/api/uploads/ref.docx"]))
    assert text is not None and "Эталонный ответ учителя." in text
    assert len(images) == 1
    assert images[0].startswith("data:image/png;base64,")


def test_resolve_reference_material_empty_urls_returns_nothing():
    text, images = asyncio.run(ai_grader._resolve_reference_material([]))
    assert text is None
    assert images == []


# ── grade_submission/grade_handwritten_submission: эталон доходит до vision ─

def test_grade_submission_sends_reference_images_with_distinct_caption(monkeypatch):
    """Картинки эталона учителя должны попасть в промпт ОТДЕЛЬНО от картинок
    студента, с явной подписью, что это не работа студента — иначе модель
    может перепутать, чью работу она оценивает."""
    monkeypatch.setenv("OPENAI_API_KEY", "test-key")
    ref_image = _png_bytes((9, 9, 9))
    data = _docx_with_image("Что должно быть в ответе.", ref_image)

    # httpx используется и для скачивания reference-файла (GET), и для
    # вызова OpenAI (POST) — единый фейковый клиент обслуживает оба.
    class _DualClient(_CapturingClient):
        def __init__(self, *a, **k):
            super().__init__(*a, **k)

        async def get(self, *a, **k):
            class _R:
                is_success = True
                status_code = 200
                content = data
                headers = {"content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
            return _R()

    monkeypatch.setattr(ai_grader.httpx, "AsyncClient", _DualClient)

    result = asyncio.run(ai_grader.grade_submission(
        text="Ответ студента текстом.",
        criteria=[{"name": "полнота", "weight": 100}],
        max_score=100,
        reference_solution_urls=["http://localhost:8000/api/uploads/ref.docx"],
    ))
    assert result["score"] == 80  # см. _CapturingClient — фиксированный ответ

    payload = _CapturingClient.last_payload
    user_msg = next(m["content"] for m in payload["messages"] if m["role"] == "user")
    assert isinstance(user_msg, list)
    captions = [p["text"] for p in user_msg if p.get("type") == "text"]
    assert any("ЭТАЛОННОГО решения учителя" in c for c in captions)
    assert any(p.get("type") == "image_url" for p in user_msg)


# ── Сквозной паритет: deadline_checker._grade_one использует ТОТ ЖЕ эталон ──

def test_deadline_checker_uses_variant_reference_same_as_manual_path(db_session, monkeypatch):
    """До фикса автопроверка по дедлайну игнорировала submission.variant_number
    и брала assignment.reference_solution_url — здесь у задания намеренно
    задан ДРУГОЙ (неверный) эталон на уровне задания, чтобы тест падал, если
    регресс вернётся."""
    monkeypatch.setenv("OPENAI_API_KEY", "test-key")
    ref_data = _docx_with_image("Эталон варианта 2.", _png_bytes((5, 5, 5)))

    class _DualClient(_CapturingClient):
        async def get(self, *a, **k):
            class _R:
                is_success = True
                status_code = 200
                content = ref_data
                headers = {"content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
            return _R()

    monkeypatch.setattr(ai_grader.httpx, "AsyncClient", _DualClient)

    teacher = make_user(db_session, role="teacher")
    _, assignment = _make_assignment(
        db_session, teacher, reference_solution_url="http://x/wrong-assignment-level.docx"
    )
    crud_classes.add_variant(db_session, assignment.id, 2, "http://localhost:8000/api/uploads/variant2.docx")

    student = make_user(db_session, role="student")
    sub = Submission(
        assignment_id=assignment.id, student_id=student.id, variant_number=2,
        text_content="Ответ студента по варианту 2.", status="submitted",
    )
    db_session.add(sub)
    db_session.commit()
    db_session.refresh(sub)

    asyncio.run(deadline_checker._grade_one(db_session, sub, assignment, org_type="university"))

    payload = _CapturingClient.last_payload
    assert payload is not None, "автопроверка не дошла до вызова модели"
    user_msg = next(m["content"] for m in payload["messages"] if m["role"] == "user")
    assert isinstance(user_msg, list)  # vision-контент — есть картинка эталона
    text_blocks = " ".join(p["text"] for p in user_msg if p.get("type") == "text")
    assert "Эталон варианта 2." in text_blocks
    assert any(p.get("type") == "image_url" for p in user_msg)

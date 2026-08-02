"""Извлечение встроенных в .docx изображений (скриншоты/сканы/фото/схемы,
вставленные ВМЕСТО текста) и колонтитулов — раньше _parse_docx видел только
doc.paragraphs/doc.tables (текстовый слой тела документа) и полностью
игнорировал картинки и header/footer. Студент не должен терять баллы за то,
что ответ вставлен как картинка, а не напечатан."""
import asyncio
import io
import zipfile

import pytest
from docx import Document
from docx.shared import Inches
from PIL import Image

from services import ai_grader


def _png_bytes(color: tuple) -> bytes:
    buf = io.BytesIO()
    Image.new("RGB", (20, 20), color=color).save(buf, format="PNG")
    return buf.getvalue()


IMAGE_A = _png_bytes((255, 0, 0))   # красный — "тело документа"
IMAGE_B = _png_bytes((0, 0, 255))   # синий — "колонтитул"


def _build_docx(with_header_image: bool = True) -> bytes:
    doc = Document()
    doc.add_heading("Задание", level=1)
    p = doc.add_paragraph("Ответ на первый вопрос — см. схему ниже.")
    p.add_run().add_picture(io.BytesIO(IMAGE_A), width=Inches(1))

    # Та же картинка ещё раз, теперь в ячейке таблицы — проверка дедупа И
    # того, что изображения в таблицах вообще извлекаются.
    table = doc.add_table(rows=1, cols=1)
    table.rows[0].cells[0].paragraphs[0].add_run().add_picture(
        io.BytesIO(IMAGE_A), width=Inches(1)
    )

    doc.add_paragraph("Второй вопрос — текстом.")

    section = doc.sections[0]
    section.header.is_linked_to_previous = False
    header_p = section.header.paragraphs[0]
    header_p.text = "ФИО: Иванов И.И., вариант 2"
    if with_header_image:
        header_p.add_run().add_picture(io.BytesIO(IMAGE_B), width=Inches(0.5))

    section.footer.is_linked_to_previous = False
    section.footer.paragraphs[0].text = "Страница 1"

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── _parse_docx: текст + колонтитулы ────────────────────────────────────────

def test_parse_docx_extracts_body_text():
    text = ai_grader._parse_docx(_build_docx())
    assert "Ответ на первый вопрос" in text
    assert "Второй вопрос — текстом." in text


def test_parse_docx_extracts_header_and_footer():
    """Регресс: doc.paragraphs не включает header/footer вообще — раньше эти
    строки терялись, даже если студент писал ФИО/вариант/ответ в колонтитуле."""
    text = ai_grader._parse_docx(_build_docx())
    assert "Иванов И.И." in text
    assert "Страница 1" in text


# ── _extract_docx_images: порядок, дедуп, таблицы, колонтитулы ──────────────

def test_extract_docx_images_finds_body_and_header_images():
    images = ai_grader._extract_docx_images(_build_docx())
    exts = [ext for _, ext in images]
    assert len(images) == 2  # IMAGE_A (дважды, дедуп) + IMAGE_B
    assert all(e == "png" for e in exts)


def test_extract_docx_images_deduplicates_identical_image_in_table():
    """IMAGE_A вставлена в тело И в ячейку таблицы — на выходе должна быть
    ровно одна копия, а не две одинаковых."""
    images = ai_grader._extract_docx_images(_build_docx())
    raw_bytes = [data for data, _ in images]
    assert raw_bytes.count(IMAGE_A) == 1


def test_extract_docx_images_orders_body_before_header():
    """Картинка из тела документа (IMAGE_A) должна идти раньше картинки из
    колонтитула (IMAGE_B) — порядок появления в документе, а не порядок
    файлов в word/media/."""
    images = ai_grader._extract_docx_images(_build_docx())
    raw_bytes = [data for data, _ in images]
    assert raw_bytes.index(IMAGE_A) < raw_bytes.index(IMAGE_B)


def test_extract_docx_images_respects_max_images_cap():
    doc = Document()
    for i in range(ai_grader.MAX_DOCX_EMBEDDED_IMAGES + 3):
        p = doc.add_paragraph()
        p.add_run().add_picture(io.BytesIO(_png_bytes((i % 255, 0, 0))), width=Inches(0.2))
    buf = io.BytesIO()
    doc.save(buf)
    images = ai_grader._extract_docx_images(buf.getvalue())
    assert len(images) == ai_grader.MAX_DOCX_EMBEDDED_IMAGES


def test_extract_docx_images_empty_for_docx_without_images():
    doc = Document()
    doc.add_paragraph("Просто текст, без картинок.")
    buf = io.BytesIO()
    doc.save(buf)
    assert ai_grader._extract_docx_images(buf.getvalue()) == []


def test_extract_docx_images_corrupted_bytes_returns_empty_not_raise():
    assert ai_grader._extract_docx_images(b"not a docx at all") == []


def test_extract_docx_images_resolves_colliding_rids_across_parts():
    """Регресс: rId — НЕ глобальный идентификатор, каждая часть пакета
    (document.xml, header1.xml, footer1.xml...) нумерует свои связи заново
    со своего собственного _rels-файла. python-docx реально генерирует
    header1.xml.rels и footer1.xml.rels, ОБА использующие "rId1" для разных
    картинок — раньше единый общий словарь {rid: путь} на весь пакет тихо
    схлопывал такие связи в одну (побеждала связь, обработанная последней),
    и одно из двух изображений либо терялось, либо резолвилось не туда."""
    image_header = _png_bytes((0, 255, 0))
    image_footer = _png_bytes((255, 255, 0))

    doc = Document()
    doc.add_paragraph("Тело без картинок — коллизия только в колонтитулах.")
    section = doc.sections[0]
    section.header.is_linked_to_previous = False
    section.header.paragraphs[0].add_run().add_picture(io.BytesIO(image_header), width=Inches(0.5))
    section.footer.is_linked_to_previous = False
    section.footer.paragraphs[0].add_run().add_picture(io.BytesIO(image_footer), width=Inches(0.5))
    buf = io.BytesIO()
    doc.save(buf)
    data = buf.getvalue()

    # Убеждаемся, что тестовый файл действительно воспроизводит коллизию
    # (иначе тест ничего не проверяет) — оба .rels независимо используют rId1.
    with zipfile.ZipFile(io.BytesIO(data)) as z:
        header_rels = z.read("word/_rels/header1.xml.rels").decode()
        footer_rels = z.read("word/_rels/footer1.xml.rels").decode()
    assert 'Id="rId1"' in header_rels and 'Id="rId1"' in footer_rels

    images = ai_grader._extract_docx_images(data)
    raw_bytes = [d for d, _ in images]
    assert len(images) == 2
    assert image_header in raw_bytes
    assert image_footer in raw_bytes
    # Правильное разрешение по частям означает и правильный порядок: header
    # закономерно раньше footer. При схлопывании в общий словарь порядок
    # ломался (последняя обработанная связь "побеждала" для всех частей).
    assert raw_bytes.index(image_header) < raw_bytes.index(image_footer)


def test_extract_docx_images_headers_ordered_before_footers():
    """Колонтитулы раньше сортировались одним списком по алфавиту имени
    файла ("footer1.xml" < "header1.xml"), из-за чего футер попадал в
    контекст раньше хедера — не соответствует читаемому порядку страницы."""
    image_header = _png_bytes((10, 20, 30))
    image_footer = _png_bytes((40, 50, 60))

    doc = Document()
    doc.add_paragraph("Текст.")
    section = doc.sections[0]
    section.header.is_linked_to_previous = False
    section.header.paragraphs[0].add_run().add_picture(io.BytesIO(image_header), width=Inches(0.3))
    section.footer.is_linked_to_previous = False
    section.footer.paragraphs[0].add_run().add_picture(io.BytesIO(image_footer), width=Inches(0.3))
    buf = io.BytesIO()
    doc.save(buf)

    images = ai_grader._extract_docx_images(buf.getvalue())
    raw_bytes = [d for d, _ in images]
    assert raw_bytes.index(image_header) < raw_bytes.index(image_footer)


# ── _image_bytes_to_data_uri ────────────────────────────────────────────────

def test_image_bytes_to_data_uri_valid_png():
    uri = ai_grader._image_bytes_to_data_uri(IMAGE_A, "png")
    assert uri.startswith("data:image/png;base64,")


def test_image_bytes_to_data_uri_rejects_unknown_extension():
    assert ai_grader._image_bytes_to_data_uri(IMAGE_A, "emf") is None


def test_image_bytes_to_data_uri_rejects_oversized_image():
    huge = b"x" * (ai_grader.MAX_DOCX_IMAGE_BYTES + 1)
    assert ai_grader._image_bytes_to_data_uri(huge, "png") is None


# ── _fetch_docx_embedded_images: скачивание + извлечение ────────────────────

class _FakeGetResp:
    def __init__(self, content: bytes):
        self.content = content
        self.is_success = True
        self.status_code = 200
        self.headers = {"content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"}


class _FakeGetClient:
    def __init__(self, *a, **k):
        pass

    async def __aenter__(self):
        return self

    async def __aexit__(self, *a):
        return False

    async def get(self, *a, **k):
        return _FakeGetResp(_build_docx())


def test_fetch_docx_embedded_images_downloads_and_extracts(monkeypatch):
    monkeypatch.setattr(ai_grader.httpx, "AsyncClient", _FakeGetClient)
    uris = asyncio.run(
        ai_grader._fetch_docx_embedded_images("http://localhost:8000/api/uploads/hw.docx")
    )
    assert len(uris) == 2
    assert all(u.startswith("data:image/png;base64,") for u in uris)


def test_fetch_docx_embedded_images_skips_non_docx_urls(monkeypatch):
    monkeypatch.setattr(ai_grader.httpx, "AsyncClient", _FakeGetClient)
    uris = asyncio.run(
        ai_grader._fetch_docx_embedded_images("http://localhost:8000/api/uploads/hw.pdf")
    )
    assert uris == []


# ── grade_submission: встроенные картинки реально доходят до промпта ────────

class _CapturingPostClient:
    last_messages = None

    def __init__(self, *a, **k):
        pass

    async def __aenter__(self):
        return self

    async def __aexit__(self, *a):
        return False

    async def post(self, url, headers=None, json=None):
        _CapturingPostClient.last_messages = json["messages"]

        class _R:
            is_success = True
            status_code = 200

            def json(self_inner):
                import json as _j
                return {
                    "choices": [{"message": {"content": _j.dumps({
                        "confidence": 90, "confidence_reasons": [], "score": 100,
                        "feedback": "ok",
                        "criteria_scores": [{"name": "полнота", "score": 100, "max": 100, "comment": "ok"}],
                    })}}],
                    "usage": {},
                }
        return _R()


def test_grade_submission_sends_embedded_images_to_vision(monkeypatch):
    monkeypatch.setenv("OPENAI_API_KEY", "test-key")
    monkeypatch.setattr(ai_grader.httpx, "AsyncClient", _CapturingPostClient)

    data_uri = ai_grader._image_bytes_to_data_uri(IMAGE_A, "png")
    result = asyncio.run(ai_grader.grade_submission(
        text="",
        criteria=[{"name": "полнота", "weight": 100}],
        max_score=100,
        embedded_image_urls=[data_uri],
    ))
    assert result["score"] == 100

    messages = _CapturingPostClient.last_messages
    system_msg = next(m["content"] for m in messages if m["role"] == "system")
    user_msg = next(m["content"] for m in messages if m["role"] == "user")
    # Модели явно сказано не занижать балл за форму подачи (картинка вместо текста).
    assert "НЕ повод снижать балл" in system_msg
    # user content стал vision-списком (текст + картинка), а не плоской строкой.
    assert isinstance(user_msg, list)
    assert any(part.get("type") == "image_url" for part in user_msg)
    image_part = next(part for part in user_msg if part.get("type") == "image_url")
    assert image_part["image_url"]["url"] == data_uri


def test_grade_submission_without_images_stays_plain_text_prompt(monkeypatch):
    """Без картинок формат промпта не должен меняться (обратная совместимость)."""
    monkeypatch.setenv("OPENAI_API_KEY", "test-key")
    monkeypatch.setattr(ai_grader.httpx, "AsyncClient", _CapturingPostClient)

    asyncio.run(ai_grader.grade_submission(
        text="Просто текстовый ответ.",
        criteria=[{"name": "полнота", "weight": 100}],
        max_score=100,
    ))
    messages = _CapturingPostClient.last_messages
    user_msg = next(m["content"] for m in messages if m["role"] == "user")
    assert isinstance(user_msg, str)
